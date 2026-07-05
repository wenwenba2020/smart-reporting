"""Word 输出适配器 — 将 report_json 渲染为 DOCX"""
from pathlib import Path
from backend.report_engine.output_engine import OutputAdapter, OutputResult


class DocxOutputAdapter(OutputAdapter):

    @property
    def format_type(self) -> str:
        return "docx"

    async def generate(
        self, report_json: dict, template_path: str | None = None,
        style_rules: dict | None = None, output_dir: str | Path = ".",
    ) -> OutputResult:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        output_dir = Path(output_dir)
        output_dir.mkdir(parents=True, exist_ok=True)

        doc = Document(template_path) if template_path else Document()

        title = doc.add_heading(report_json.get("title", "报告"), level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        abstract = report_json.get("abstract", "")
        if abstract:
            p = doc.add_paragraph(abstract)
            try:
                p.style = doc.styles['Subtitle']
            except KeyError:
                pass

        for section in report_json.get("sections", []):
            doc.add_heading(section.get("heading", ""), level=1)
            doc.add_paragraph(section.get("body", ""))

            for table_data in section.get("tables", []):
                doc.add_paragraph(table_data.get("caption", ""))
                headers = table_data.get("headers", [])
                rows = table_data.get("rows", [])
                if headers and rows:
                    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
                    table.style = 'Light Grid Accent 1'
                    for j, h in enumerate(headers):
                        table.rows[0].cells[j].text = str(h)
                    for i, row in enumerate(rows):
                        for j, cell in enumerate(row):
                            table.rows[i + 1].cells[j].text = str(cell)

        file_path = output_dir / "report.docx"
        doc.save(str(file_path))

        return OutputResult(
            file_path=str(file_path), file_type="docx",
            mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            file_size=file_path.stat().st_size,
        )
