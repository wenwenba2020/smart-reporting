"""W2-1: 文档解析器 — PDF/DOCX/URL → Markdown"""
import re
from pathlib import Path

import fitz  # PyMuPDF


# Max pages before chunking
CHUNK_THRESHOLD = 50
CHUNK_SIZE = 20  # pages per chunk


def _page_to_markdown(page) -> str:
    """Convert a single PDF page to Markdown using structured extraction."""
    blocks = page.get_text("dict")["blocks"]
    lines: list[str] = []

    for block in blocks:
        if block["type"] != 0:  # skip images
            continue
        for line in block["lines"]:
            text = "".join(span["text"] for span in line["spans"]).strip()
            if not text:
                continue
            max_size = max(span["size"] for span in line["spans"])
            if max_size >= 20:
                lines.append(f"## {text}")
            elif max_size >= 16:
                lines.append(f"### {text}")
            else:
                lines.append(text)

    return "\n".join(lines)


def pdf_to_markdown(pdf_path: str | Path) -> str:
    """Extract text from PDF and convert to structured Markdown."""
    doc = fitz.open(str(pdf_path))
    chunks: list[str] = []

    for i in range(doc.page_count):
        page_md = _page_to_markdown(doc[i])
        if page_md:
            chunks.append(f"<!-- Page {i + 1} -->\n{page_md}")

    doc.close()
    return "\n\n".join(chunks)


def pdf_to_markdown_chunked(pdf_path: str | Path) -> list[str]:
    """For long PDFs (>50 pages), return a list of markdown chunks."""
    doc = fitz.open(str(pdf_path))
    pages = doc.page_count

    if pages <= CHUNK_THRESHOLD:
        result = []
        chunk_lines: list[str] = []
        for i in range(pages):
            page_md = _page_to_markdown(doc[i])
            if page_md:
                chunk_lines.append(f"<!-- Page {i + 1} -->\n{page_md}")
        doc.close()
        return ["\n\n".join(chunk_lines)] if chunk_lines else [""]

    # Split into chunks, using same structured extraction
    chunks: list[str] = []
    for start in range(0, pages, CHUNK_SIZE):
        end = min(start + CHUNK_SIZE, pages)
        chunk_lines = []
        for i in range(start, end):
            page_md = _page_to_markdown(doc[i])
            if page_md:
                chunk_lines.append(f"<!-- Page {i + 1} -->\n{page_md}")
        if chunk_lines:
            chunks.append("\n\n".join(chunk_lines))
    doc.close()
    return chunks


def docx_to_markdown(docx_path: str | Path) -> str:
    """Extract text from DOCX and convert to Markdown."""
    from docx import Document

    doc = Document(str(docx_path))
    lines: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        style_name = (para.style.name or "").lower()
        if "heading 1" in style_name:
            lines.append(f"# {text}")
        elif "heading 2" in style_name:
            lines.append(f"## {text}")
        elif "heading 3" in style_name:
            lines.append(f"### {text}")
        elif "list" in style_name:
            lines.append(f"- {text}")
        else:
            lines.append(text)

    # Process tables
    for table in doc.tables:
        if not table.rows:
            continue
        header = [cell.text.strip() for cell in table.rows[0].cells]
        lines.append("| " + " | ".join(header) + " |")
        lines.append("| " + " | ".join(["---"] * len(header)) + " |")
        for row in table.rows[1:]:
            cells = [cell.text.strip() for cell in row.cells]
            lines.append("| " + " | ".join(cells) + " |")
        lines.append("")

    return "\n\n".join(lines)


def url_to_markdown(url: str) -> str:
    """Fetch URL and convert to Markdown. Simple HTML stripping for MVP."""
    import httpx

    resp = httpx.get(
        url,
        follow_redirects=True,
        timeout=30,
        headers={"User-Agent": "PPTAgent/1.0"},
    )
    resp.raise_for_status()
    html = resp.text

    # Simple HTML → text extraction (MVP; Crawl4AI can replace later)
    text = re.sub(r"<script[^>]*>.*?</script>", "", html, flags=re.DOTALL)
    text = re.sub(r"<style[^>]*>.*?</style>", "", text, flags=re.DOTALL)
    text = re.sub(r"<[^>]+>", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"[ \t]+", " ", text)
    return text.strip()


def parse_document(file_path: str | Path) -> str:
    """Auto-detect format and parse to Markdown. Uses chunked path for long PDFs."""
    path = Path(file_path)
    suffix = path.suffix.lower()

    if suffix == ".pdf":
        chunks = pdf_to_markdown_chunked(path)
        return "\n\n".join(chunks)
    elif suffix == ".docx":
        return docx_to_markdown(path)
    elif suffix in (".md", ".txt"):
        return path.read_text(encoding="utf-8")
    else:
        raise ValueError(f"Unsupported file format: {suffix}. Supported: .pdf, .docx, .md, .txt")
