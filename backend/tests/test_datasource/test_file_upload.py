import pytest
import tempfile
import os
from pathlib import Path

from backend.core.datasource.file_upload import FileUploadAdapter


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _write_temp(content: str, suffix: str) -> str:
    """Write *content* to a named temp file and return the path."""
    fd, path = tempfile.mkstemp(suffix=suffix)
    with os.fdopen(fd, "w", encoding="utf-8") as fh:
        fh.write(content)
    return path


def _write_temp_bytes(data: bytes, suffix: str) -> str:
    """Write binary *data* to a named temp file and return the path."""
    fd, path = tempfile.mkstemp(suffix=suffix)
    with os.fdopen(fd, "wb") as fh:
        fh.write(data)
    return path


# ---------------------------------------------------------------------------
# Tests
# ---------------------------------------------------------------------------


@pytest.mark.asyncio
async def test_file_upload_txt():
    adapter = FileUploadAdapter()
    test_file = _write_temp("Q3营收5200万元\n利润增长15%", ".txt")
    try:
        doc = await adapter.parse(test_file, "test_sample.txt")
        assert doc.source_type == "file_upload"
        assert "5200万" in doc.content
        assert doc.title == "test_sample.txt"
    finally:
        os.unlink(test_file)


@pytest.mark.asyncio
async def test_file_upload_md():
    adapter = FileUploadAdapter()
    content = "# 报告\n\n正文内容"
    test_file = _write_temp(content, ".md")
    try:
        doc = await adapter.parse(test_file, "report.md")
        assert doc.source_type == "file_upload"
        assert "# 报告" in doc.content
    finally:
        os.unlink(test_file)


@pytest.mark.asyncio
async def test_file_upload_csv():
    adapter = FileUploadAdapter()
    content = "Name,Value,Date\nA,100,2024-01-01\nB,200,2024-02-01"
    test_file = _write_temp(content, ".csv")
    try:
        doc = await adapter.parse(test_file, "data.csv")
        assert "Name" in doc.content
        assert "A" in doc.content
        assert "---" in doc.content  # markdown table separator
    finally:
        os.unlink(test_file)


@pytest.mark.asyncio
async def test_file_upload_docx():
    """Create a minimal .docx with python-docx and verify heading detection."""
    from docx import Document

    adapter = FileUploadAdapter()

    fd, path = tempfile.mkstemp(suffix=".docx")
    os.close(fd)  # python-docx writes by path, so we close the fd

    try:
        doc = Document()
        doc.add_heading("主标题", level=1)
        doc.add_paragraph("This is a normal paragraph.")
        doc.add_heading("子标题", level=2)
        doc.add_paragraph("Another paragraph under sub-heading.")
        doc.save(path)

        result = await adapter.parse(path, "test.docx")
        assert "# 主标题" in result.content
        assert "## 子标题" in result.content
        assert "normal paragraph" in result.content
    finally:
        os.unlink(path)


@pytest.mark.asyncio
async def test_file_upload_pdf():
    """Create a minimal PDF with PyMuPDF and verify text extraction."""
    import fitz

    adapter = FileUploadAdapter()

    fd, path = tempfile.mkstemp(suffix=".pdf")
    os.close(fd)

    try:
        pdf = fitz.open()
        pdf.insert_page(-1, text="Q3 Revenue: 52,000,000 CNY\nGrowth: 15%", fontsize=12)
        pdf.save(path)
        pdf.close()

        result = await adapter.parse(path, "report.pdf")
        assert "Q3 Revenue" in result.content or "52,000,000" in result.content
        assert "Page 1" in result.content
    finally:
        os.unlink(path)


@pytest.mark.asyncio
async def test_file_upload_xlsx():
    """Create a minimal .xlsx with openpyxl and verify extraction."""
    from openpyxl import Workbook

    adapter = FileUploadAdapter()

    fd, path = tempfile.mkstemp(suffix=".xlsx")
    os.close(fd)

    try:
        wb = Workbook()
        ws = wb.active
        ws.title = "Sales"
        ws.append(["Product", "Revenue"])
        ws.append(["Widget", "5000"])
        ws.append(["Gadget", "3000"])
        wb.save(path)
        wb.close()

        result = await adapter.parse(path, "sales.xlsx")
        assert "Sheet: Sales" in result.content
        assert "Widget" in result.content
        assert "5000" in result.content
    finally:
        os.unlink(path)


@pytest.mark.asyncio
async def test_file_upload_image():
    """Image files produce placeholder text."""
    from PIL import Image

    adapter = FileUploadAdapter()

    fd, path = tempfile.mkstemp(suffix=".png")
    os.close(fd)

    try:
        img = Image.new("RGB", (100, 200), color="red")
        img.save(path)

        result = await adapter.parse(path, "photo.png")
        assert "[Image:" in result.content
        assert "100" in result.content
    finally:
        os.unlink(path)


@pytest.mark.parametrize(
    "filename,expected",
    [
        ("report.docx", True),
        ("data.xlsx", True),
        ("notes.txt", True),
        ("readme.md", True),
        ("export.csv", True),
        ("scan.pdf", True),
        ("photo.png", True),
        ("photo.jpg", True),
        ("photo.jpeg", True),
        ("unknown.xyz", False),
        ("script.py", False),
        ("noext", False),
    ],
)
def test_file_upload_supports(filename, expected):
    adapter = FileUploadAdapter()
    assert adapter.supports(filename) is expected


@pytest.mark.asyncio
async def test_file_upload_metadata_merge():
    adapter = FileUploadAdapter()
    test_file = _write_temp("hello", ".txt")
    try:
        doc = await adapter.parse(
            test_file, "hello.txt", metadata={"user_id": "u1"}
        )
        assert doc.metadata["user_id"] == "u1"
    finally:
        os.unlink(test_file)
