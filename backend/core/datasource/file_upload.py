"""
FileUploadAdapter — parses uploaded files into SourceDocument instances.

Supported formats:
    .txt / .md   → plain-text read
    .csv          → csv read
    .docx         → python-docx paragraph extraction with heading detection
    .pdf          → PyMuPDF (fitz) text per page
    .xlsx         → openpyxl all-sheet extraction
    .png / .jpg   → image placeholder (metadata note)
"""

from __future__ import annotations

import csv
import uuid
from pathlib import Path
from typing import Optional

from backend.core.datasource.base import DataSourceAdapter, register_adapter
from backend.core.models import SourceDocument

# ---------------------------------------------------------------------------
# Supported extensions
# ---------------------------------------------------------------------------

_TEXT_EXTS = {".txt", ".md"}
_IMAGE_EXTS = {".png", ".jpg", ".jpeg"}

_SUPPORTED_EXTS = _TEXT_EXTS | _IMAGE_EXTS | {".csv", ".docx", ".pdf", ".xlsx"}


@register_adapter
class FileUploadAdapter(DataSourceAdapter):
    source_type = "file_upload"

    # ------------------------------------------------------------------
    # supports
    # ------------------------------------------------------------------

    def supports(self, filename: str) -> bool:
        ext = Path(filename).suffix.lower()
        return ext in _SUPPORTED_EXTS

    # ------------------------------------------------------------------
    # parse — dispatcher
    # ------------------------------------------------------------------

    async def parse(
        self,
        file_path: str,
        filename: str,
        metadata: Optional[dict] = None,
    ) -> SourceDocument:
        ext = Path(filename).suffix.lower()

        if ext in _TEXT_EXTS:
            content = await self._parse_text(file_path)
        elif ext == ".csv":
            content = await self._parse_csv(file_path)
        elif ext == ".docx":
            content = await self._parse_docx(file_path)
        elif ext == ".pdf":
            content = await self._parse_pdf(file_path)
        elif ext == ".xlsx":
            content = await self._parse_excel(file_path)
        elif ext in _IMAGE_EXTS:
            content = await self._parse_image(file_path, filename)
        else:
            raise ValueError(f"Unsupported file extension: {ext}")

        return SourceDocument(
            id=str(uuid.uuid4()),
            title=filename,
            content=content,
            source_type=self.source_type,
            metadata=metadata or {},
        )

    # ------------------------------------------------------------------
    # Format-specific parsers
    # ------------------------------------------------------------------

    async def _parse_text(self, file_path: str) -> str:
        with open(file_path, "r", encoding="utf-8", errors="replace") as fh:
            return fh.read()

    async def _parse_csv(self, file_path: str) -> str:
        lines: list[str] = []
        with open(file_path, "r", encoding="utf-8", errors="replace") as fh:
            reader = csv.reader(fh)
            headers = next(reader, None)
            if headers:
                lines.append(" | ".join(headers))
                lines.append(" | ".join(["---"] * len(headers)))
            for row in reader:
                lines.append(" | ".join(row))
        return "\n".join(lines)

    async def _parse_docx(self, file_path: str) -> str:
        from docx import Document  # type: ignore[import-untyped]

        doc = Document(file_path)
        parts: list[str] = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            style_name = para.style.name if para.style else ""

            # Detect heading styles
            if any(
                keyword in style_name
                for keyword in ("Heading", "heading", "标题")
            ):
                # Try to extract heading level
                level = 1
                for char in style_name:
                    if char.isdigit():
                        level = int(char)
                        break
                prefix = "#" * min(level, 6)
                parts.append(f"{prefix} {text}")
            else:
                parts.append(text)

        return "\n\n".join(parts)

    async def _parse_pdf(self, file_path: str) -> str:
        import fitz  # type: ignore[import-untyped]

        doc = fitz.open(file_path)
        parts: list[str] = []
        for i, page in enumerate(doc, start=1):
            text = page.get_text("text")
            if text.strip():
                parts.append(f"## Page {i}\n\n{text.strip()}")
        doc.close()
        return "\n\n".join(parts)

    async def _parse_excel(self, file_path: str) -> str:
        from openpyxl import load_workbook  # type: ignore[import-untyped]

        wb = load_workbook(file_path, data_only=True)
        parts: list[str] = []

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            parts.append(f"## Sheet: {sheet_name}\n")

            # Text representation — all rows as pipe-separated values
            rows: list[str] = []
            for row in ws.iter_rows(values_only=True):
                row_vals = [
                    str(cell) if cell is not None else "" for cell in row
                ]
                rows.append(" | ".join(row_vals))
            if rows:
                parts.append("\n".join(rows))

            # Column summaries
            parts.append(f"\n### {sheet_name} - Column Summaries\n")
            for col_cells in ws.iter_cols(values_only=True):
                col_vals = [
                    str(c) for c in col_cells if c is not None
                ]
                if col_vals:
                    parts.append(f"- {', '.join(col_vals[:20])}")

        wb.close()
        return "\n\n".join(parts)

    async def _parse_image(self, file_path: str, filename: str) -> str:
        """Return placeholder text for images — actual OCR is handled elsewhere."""
        from PIL import Image  # type: ignore[import-untyped]

        img = Image.open(file_path)
        w, h = img.size
        img.close()
        return (
            f"[Image: {filename}]\n"
            f"Dimensions: {w}x{h}\n"
            f"Note: Image content requires OCR processing for text extraction."
        )
