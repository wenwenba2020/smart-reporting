import os
import re
from backend.core.output.base import BaseExporter
from backend.core.models import ExportResult, StructuredReport, ReportSection


_TABLE_ROW_RE = re.compile(r"^\|(.+)\|$")
_TABLE_SEP_RE = re.compile(r"^\|[\s\-:]+\|$")


class DocxExporter(BaseExporter):
    format = "docx"

    async def export(
        self,
        report: StructuredReport,
        output_dir: str = "./data/exports",
    ) -> ExportResult:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        os.makedirs(output_dir, exist_ok=True)

        doc = Document()

        # -- Default font style --
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Arial"
        font.size = Pt(11)

        # -- Title --
        title_para = doc.add_heading(report.meta.title, level=1)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # -- Meta info --
        meta = report.meta
        meta_parts = []
        if meta.author:
            meta_parts.append(f"作者: {meta.author}")
        if meta.created_at:
            meta_parts.append(f"日期: {meta.created_at}")
        if meta.template_id:
            meta_parts.append(f"模板: {meta.template_id}")
        if meta.tags:
            meta_parts.append(f"标签: {', '.join(meta.tags)}")
        if meta_parts:
            meta_line = " | ".join(meta_parts)
            meta_para = doc.add_paragraph(meta_line)
            meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in meta_para.runs:
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(128, 128, 128)

        # -- Separator --
        doc.add_paragraph("─" * 60)

        # -- Sections --
        for section in report.sections:
            heading_text = section.section_def.title or section.section_def.heading
            if heading_text:
                doc.add_heading(heading_text, level=2)

            content = section.markdown_content or ""
            self._render_content(doc, content)

        # -- If there are no sections but there is full_markdown, render that --
        if not report.sections and report.full_markdown:
            self._render_content(doc, report.full_markdown)

        # -- Save --
        file_name = f"{report.meta.id}.docx"
        file_path = os.path.join(output_dir, file_name)
        doc.save(file_path)
        file_size = os.path.getsize(file_path)

        return ExportResult(
            file_path=file_path,
            format=self.format,
            file_size_bytes=file_size,
        )

    def _render_content(self, doc, content: str) -> None:
        """Parse markdown-like content and add it to the docx document."""
        from docx.shared import Pt, RGBColor

        lines = content.split("\n")
        table_buffer: list[str] = []

        for line in lines:
            # Check if we're in table mode
            if table_buffer:
                table_buffer.append(line)
                continue

            # Detect table start: first row of a pipe table
            stripped = line.strip()

            # Heading H3 (###)
            if stripped.startswith("### "):
                doc.add_heading(stripped[4:], level=3)
                continue

            # Heading H2 (##)
            if stripped.startswith("## "):
                doc.add_heading(stripped[3:], level=2)
                continue

            # Table row
            if _TABLE_ROW_RE.match(stripped):
                table_buffer.append(stripped)
                continue

            # Horizontal rule
            if stripped in ("---", "***", "___"):
                doc.add_paragraph("─" * 40)
                continue

            # Unordered list item
            if stripped.startswith(("- ", "* ", "+ ")):
                text = stripped[2:].strip()
                para = doc.add_paragraph(text, style="List Bullet")
                self._style_run(para, text)
                continue

            # Ordered list item
            if re.match(r"^\d+\.\s", stripped):
                text = re.sub(r"^\d+\.\s", "", stripped).strip()
                para = doc.add_paragraph(text, style="List Number")
                self._style_run(para, text)
                continue

            # Empty line
            if not stripped:
                # Flush table buffer if any
                if table_buffer:
                    self._flush_table(doc, table_buffer)
                    table_buffer = []
                continue

            # Regular paragraph
            if stripped:
                para = doc.add_paragraph(stripped)
                self._style_run(para, stripped)

        # Flush any remaining table buffer
        if table_buffer:
            self._flush_table(doc, table_buffer)

    def _flush_table(self, doc, rows: list[str]) -> None:
        """Convert buffered pipe-table rows into a docx Table."""
        if not rows:
            return

        from docx.shared import Pt, RGBColor, Inches
        from docx.oxml.ns import qn

        # Parse header row and detect if there's a separator
        header_cells = [c.strip() for c in rows[0].strip("|").split("|")]
        data_start = 1
        if len(rows) > 1 and _TABLE_SEP_RE.match(rows[1].strip()):
            data_start = 2

        data_rows = []
        for row in rows[data_start:]:
            data_rows.append([c.strip() for c in row.strip("|").split("|")])

        num_cols = len(header_cells)
        table = doc.add_table(rows=1 + len(data_rows), cols=num_cols)
        table.style = "Light Grid Accent 1"

        # Header
        for j, cell_text in enumerate(header_cells):
            cell = table.rows[0].cells[j]
            cell.text = cell_text
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.bold = True
                    run.font.size = Pt(10)

        # Data rows
        for i, row_data in enumerate(data_rows):
            for j, cell_text in enumerate(row_data):
                if j < num_cols:
                    cell = table.rows[i + 1].cells[j]
                    cell.text = cell_text
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.font.size = Pt(10)

        doc.add_paragraph()  # spacing after table

    @staticmethod
    def _style_run(para, text: str) -> None:
        """Apply inline bold (**bold**) formatting to paragraph runs."""
        from docx.shared import Pt

        bold_re = re.compile(r"\*\*(.+?)\*\*")
        if bold_re.search(text):
            parts = bold_re.split(text)
            for idx, part in enumerate(parts):
                if idx % 2 == 0:
                    run = para.add_run(part)
                else:
                    run = para.add_run(part)
                    run.bold = True
        else:
            para.clear()
            run = para.add_run(text)
            run.font.size = Pt(11)
